import os
import re
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import streamlit as st
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document as LangchainDocument
import docx
from docx.opc.exceptions import PackageNotFoundError as DocxPackageNotFoundError
import pdfplumber
from .config import PDF_STORAGE_PATH, CONTEXT_PDF_STORAGE_PATH
from .logger_config import get_logger
from .model_loader import get_language_model
from .generation import classify_chunk
from .config import (
    MODEL_CACHE_DIR,
    HF_LOCAL_FILES_ONLY,
    TOKENIZER_LOCAL_PATH,
    TOKENIZER_MODEL_NAME,
)
from .heading_cleanup import (
    extract_toc_lines, build_toc_set,                      # ToC
    clean_heading_nodes,                                   # cleanup
    build_toc_index, expand_headings_with_parents,         # parents
)

# Docling imports
from docling.chunking import HybridChunker
from docling_core.transforms.chunker.tokenizer.huggingface import HuggingFaceTokenizer
from docling_core.types import DoclingDocument
from docling_core.types.doc import DocItemLabel
from transformers import AutoTokenizer
from urllib.parse import unquote, urlparse
from docling_core.transforms.chunker.hierarchical_chunker import (
    ChunkingDocSerializer, ChunkingSerializerProvider,
)
from docling_core.transforms.serializer.markdown import MarkdownTableSerializer

logger = get_logger(__name__)

# Parallel Ollama calls during chunk classification; keep modest to avoid overloading.
_CLASSIFY_MAX_WORKERS = 4

class MDTableSerializerProvider(ChunkingSerializerProvider):
    def get_serializer(self, doc):
        return ChunkingDocSerializer(doc=doc, table_serializer=MarkdownTableSerializer())
    
def _resolve_max_tokens(tokenizer) -> int:
    """
    Determine a safe max token length for Docling's HF tokenizer without
    hitting the Hub (offline-safe). Falls back to 512 if unavailable.
    """
    try:
        max_len = getattr(tokenizer, "model_max_length", None)
        if max_len and max_len != float("inf"):
            # cap to a sane value to avoid giant defaults (e.g., 1e12)
            return int(min(max_len, 1024))
    except Exception:
        pass
    return 512


def _ensure_local_tokenizer_assets() -> bool:
    """
    Verify required tokenizer files exist when running in offline mode.
    """
    if not HF_LOCAL_FILES_ONLY:
        return True

    required_files = [
        "config.json",
        "tokenizer.json",
        "sentence_bert_config.json",
    ]
    missing = [f for f in required_files if not (TOKENIZER_LOCAL_PATH / f).exists()]
    if missing:
        user_message = (
            "Offline mode requires cached tokenizer assets. "
            f"Missing files at {TOKENIZER_LOCAL_PATH}: {', '.join(missing)}. "
            "Run `python pre_download_model.py` while online to populate the cache."
        )
        logger.error(user_message)
        st.error(user_message)
        return False
    return True


def _load_docling_tokenizer():
    """
    Load the tokenizer required by Docling's HybridChunker from a local cache.
    """
    if not _ensure_local_tokenizer_assets():
        return None

    tokenizer_source = TOKENIZER_LOCAL_PATH if TOKENIZER_LOCAL_PATH.exists() else TOKENIZER_MODEL_NAME
    try:
        logger.info(f"Loading Docling tokenizer from: {tokenizer_source}")
        return AutoTokenizer.from_pretrained(
            tokenizer_source,
            cache_dir=str(MODEL_CACHE_DIR),
            local_files_only=HF_LOCAL_FILES_ONLY,
        )
    except Exception as exc:
        user_message = (
            "Failed to load the Docling tokenizer from the local cache. "
            "Run `python pre_download_model.py` while online to populate the cache."
        )
        logger.exception(f"{user_message} Details: {exc}")
        st.error(user_message)
        return None


def _split_text_blocks(text: str) -> list[str]:
    """Split raw text into paragraph-like blocks for chunking."""
    return [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]


def _build_docling_document_from_text(text_blocks: list[str], document_name: str) -> DoclingDocument:
    """
    Build a minimal DoclingDocument from extracted text blocks. Used as a
    fallback when native Docling conversion is unavailable (e.g. the source file
    is not on disk). Note: this produces TEXT-only items with no SECTION_HEADER
    structure, so chunks built from it carry no heading hierarchy.
    """
    doc = DoclingDocument(name=document_name or "document")
    for block in text_blocks:
        for paragraph in _split_text_blocks(block):
            doc.add_text(label=DocItemLabel.TEXT, text=paragraph)
    return doc


# Cache Docling converters so layout models are loaded at most once per process.
_DOCLING_CONVERTERS: dict[bool, object] = {}


def _get_docling_converter(is_pdf: bool):
    """
    Return a cached Docling DocumentConverter. For PDFs, OCR and table-structure
    detection are disabled: the structural (heading) information needed for the
    requirements hierarchy is recovered from the layout model alone, while
    skipping the expensive OCR pass that previously made processing hang.
    """
    if is_pdf not in _DOCLING_CONVERTERS:
        from docling.document_converter import DocumentConverter

        if is_pdf:
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            from docling.document_converter import PdfFormatOption

            pdf_options = PdfPipelineOptions()
            pdf_options.do_ocr = True
            pdf_options.do_table_structure = True
            _DOCLING_CONVERTERS[is_pdf] = DocumentConverter(
                format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options)}
            )
        else:
            _DOCLING_CONVERTERS[is_pdf] = DocumentConverter()
    return _DOCLING_CONVERTERS[is_pdf]


def _convert_with_docling(full_path: str, is_pdf: bool):
    """
    Run Docling's native conversion to obtain a DoclingDocument that preserves
    SECTION_HEADER structure. Returns None on failure so callers can fall back
    to the text-only document builder.
    """
    try:
        converter = _get_docling_converter(is_pdf)
        return converter.convert(source=full_path).document
    except Exception as exc:
        logger.warning(
            f"Native Docling conversion failed for '{full_path}': {exc}. "
            "Falling back to text-only document (headings will be unavailable)."
        )
        return None


def _normalize_source_path(source_path: str) -> str:
    """Normalize file paths that may be URL-encoded or file:// URLs."""
    cleaned = unquote(source_path)
    if cleaned.startswith("file://"):
        parsed = urlparse(cleaned)
        if parsed.path:
            cleaned = parsed.path
    if "?" in cleaned:
        cleaned = cleaned.split("?", 1)[0]
    if "#" in cleaned:
        cleaned = cleaned.split("#", 1)[0]
    return cleaned


def _is_pdf_source(
    source_path: str,
    source_docs: list[LangchainDocument],
    normalized_source: str | None = None,
) -> bool:
    normalized = normalized_source or _normalize_source_path(source_path)
    if Path(normalized).suffix.lower() == ".pdf":
        return True
    # Fallback: PDFPlumberLoader adds page metadata; use that to detect PDFs.
    return any("page" in doc.metadata or "total_pages" in doc.metadata for doc in source_docs)


def save_uploaded_file(uploaded_file, storage_path=PDF_STORAGE_PATH):
    # Create the storage directory if it doesn't exist
    os.makedirs(storage_path, exist_ok=True)
    file_path = os.path.join(storage_path, uploaded_file.name)
    try:
        with open(file_path, "wb") as file:
            file.write(uploaded_file.getbuffer())
        logger.info(f"File '{uploaded_file.name}' saved to '{file_path}'.")
        return file_path
    except IOError as e:
        user_message = f"Failed to save uploaded file '{uploaded_file.name}'. An I/O error occurred: {e.strerror}."
        logger.error(f"{user_message} Please check permissions and disk space.")
        st.error(user_message)
        return None
    except Exception as e:
        user_message = f"An unexpected error occurred while saving '{uploaded_file.name}'."
        logger.exception(f"{user_message} Details: {e}")
        st.error(f"{user_message} Check logs for details.")
        return None

def load_document(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)

    try:
        if file_extension == ".pdf":
            try:
                logger.debug(f"Loading PDF: {file_name}")
                document_loader = PDFPlumberLoader(file_path)
                docs = document_loader.load()
                logger.info(f"Successfully loaded PDF: {file_name}, {len(docs)} pages/documents.")
                return docs
            except pdfplumber.exceptions.PDFSyntaxError as pdf_err:
                user_message = f"Failed to load PDF '{file_name}': The file may be corrupted or not a valid PDF."
                logger.error(f"{user_message} Details: {pdf_err}")
                st.error(user_message)
                return []
            except Exception as e:
                user_message = f"Failed to load PDF '{file_name}': An unexpected error occurred during PDF processing."
                logger.exception(f"{user_message} Details: {e}")
                st.error(f"{user_message} Check logs for details.")
                return []
        elif file_extension == ".docx":
            try:
                logger.debug(f"Loading DOCX: {file_name}")
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                if not full_text.strip():
                    logger.warning(f"DOCX file '{file_name}' is empty or contains no text.")
                    st.warning(f"DOCX file '{file_name}' appears to be empty or contains no text.")
                    return []
                logger.info(f"Successfully loaded DOCX: {file_name}")
                return [LangchainDocument(page_content=full_text, metadata={"source": file_name})]
            except DocxPackageNotFoundError as e:
                user_message = f"Failed to load DOCX '{file_name}': The file appears to be corrupted or not a valid DOCX file."
                logger.error(f"{user_message} Error: {e}")
                st.error(user_message)
                return []
            except Exception as e:
                user_message = f"Failed to load DOCX '{file_name}': An unexpected error occurred."
                logger.exception(f"{user_message} Details: {e}")
                st.error(f"{user_message} Check logs for details.")
                return []
        elif file_extension == ".txt":
            try:
                logger.debug(f"Loading TXT: {file_name}")
                with open(file_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
                if not full_text.strip():
                    logger.warning(f"Text file '{file_name}' is empty.")
                    st.warning(f"Text file '{file_name}' appears to be empty.")
                    return []
                logger.info(f"Successfully loaded TXT: {file_name}")
                return [LangchainDocument(page_content=full_text, metadata={"source": file_name})]
            except UnicodeDecodeError as unicode_err:
                user_message = f"Failed to load TXT file '{file_name}': The file is not UTF-8 encoded."
                logger.error(f"{user_message} Details: {unicode_err}")
                st.error(user_message)
                return []
            except IOError as io_err:
                user_message = f"Failed to load TXT file '{file_name}': An I/O error occurred. {io_err.strerror}."
                logger.error(f"{user_message} Details: {io_err}")
                st.error(user_message)
                return []
            except Exception as e:
                user_message = f"Failed to load TXT file '{file_name}': An unexpected error occurred."
                logger.exception(f"{user_message} Details: {e}")
                st.error(f"{user_message} Check logs for details.")
                return []
        else:
            user_message = f"Unsupported file type: '{file_extension}' for file '{file_name}'."
            logger.warning(user_message)
            st.error(user_message)
            return []
    except MemoryError as mem_err:
        user_message = f"Failed to load document '{file_name}': The file is too large to process with available memory."
        logger.error(f"{user_message} Details: {mem_err}")
        st.error(user_message)
        return []
    except Exception as e:
        user_message = f"An unexpected error occurred while attempting to load '{file_name}'."
        logger.exception(f"{user_message} Details: {e}")
        st.error(f"{user_message} Check logs for details.")
        return []


def _iter_docx_blocks(document):
    """Yield paragraphs and tables of a python-docx Document in body order."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def format_verification_context(file_path, include_tables: bool = False) -> str:
    """Read a .docx into clean plain text for use as prompt context.

    Emits each non-empty paragraph (headings and body alike) as its own line,
    preserving in-paragraph line breaks, in document order. This deliberately
    bypasses the Docling chunk pipeline, whose output for this document carries
    injected heading numbers, serialized tables, bullet markers, heading
    breadcrumbs and chunk separators. Tables are skipped by default; set
    include_tables=True to render them as 'first_cell: rest' rows.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document = docx.Document(str(file_path))
    except DocxPackageNotFoundError:
        logger.error("Verification context file not found or not a .docx: %s", file_path)
        return ""
    except Exception as exc:  # noqa: BLE001 - context is optional; never break generation
        logger.exception("Failed to read verification context '%s': %s", file_path, exc)
        return ""

    lines: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        elif include_tables and isinstance(block, Table):
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rest = " | ".join(cells[1:])
                    lines.append(cells[0] + ((": " + rest) if rest else ""))

    if not lines:
        logger.warning("Verification context produced no text: %s", file_path)
    return "\n".join(lines)


def _table_text_from_cells(table_item) -> str:
    """Render a TableItem from its raw cells, bypassing export_to_markdown()/
    export_to_text(). Both exporters share a known Docling bug
    (docling-project/docling#1575) where document-level export can silently
    skip or reorder items, which in practice drops whole table sections
    (single-row layout tables in particular). Reading table_cells directly is
    unaffected by that traversal and reliably recovers the missing rows. This
    is a safety net alongside MDTableSerializerProvider above -- markdown
    serialization fixes most empty-table cases, but a table can still
    serialize to nothing (or its whole section be skipped by the chunker) in
    which case this rebuild step recovers it independently of the chunker.
    """
    try:
        cells = table_item.data.table_cells
    except Exception:
        return ""
    if not cells:
        return ""
    rows: dict[int, dict[int, str]] = {}
    for c in cells:
        rows.setdefault(c.start_row_offset_idx, {})[c.start_col_offset_idx] = (c.text or "").strip()
    lines = []
    for r in sorted(rows):
        row = rows[r]
        line = " | ".join(row[k] for k in sorted(row) if row.get(k))
        if line:
            lines.append(line)
    return "\n".join(lines)


def _build_heading_ancestry(seq) -> dict:
    """Map each SECTION_HEADER's position in `seq` to its full ancestor chain
    (root -> leaf), reconstructed from heading levels.

    Mirrors what the HybridChunker does internally when it fills
    meta.headings for a normal chunk (it keeps a level-based stack as it
    walks the document). Recovered sections never go through the chunker, so
    this rebuilds the same ancestry independently, from the same item
    sequence, so recovered chunks get a heading chain consistent with
    everything else.
    """
    stack: list[tuple[int, str]] = []   # (level, text), root first
    ancestry: dict[int, list[str]] = {}
    for i, (item, _) in enumerate(seq):
        if item.label != DocItemLabel.SECTION_HEADER:
            continue
        level = getattr(item, "level", None) or 1
        text = item.text.strip()
        while stack and stack[-1][0] >= level:
            stack.pop()
        stack.append((level, text))
        ancestry[i] = [t for _, t in stack]
    return ancestry


def _recover_missing_table_sections(dl_doc, chunks, seq) -> list[tuple[str, str, int]]:
    """(heading, text, position) for any SECTION_HEADER whose body produced
    no chunk. `position` is the header's index in `seq` (iterate_items()),
    used to reinsert the recovered chunk at the correct place in document
    order and to look it up in _build_heading_ancestry.

    The chunker's table serialization can render a layout table (no real
    header row) as empty text; when that happens the whole section is
    skipped and its heading never appears in any chunk's meta.headings. This
    walks the same item sequence directly and rebuilds the section body from
    raw table cells and paragraph text, so the content is recovered
    regardless of how the chunker or exporters treated it.
    """
    present = {h.strip() for c in chunks for h in (c.meta.headings or [])}
    recovered = []
    for i, (item, _) in enumerate(seq):
        if item.label != DocItemLabel.SECTION_HEADER:
            continue
        heading = item.text.strip()
        if heading in present:
            continue
        parts, j = [], i + 1
        while j < len(seq) and seq[j][0].label != DocItemLabel.SECTION_HEADER:
            child = seq[j][0]
            if child.label == DocItemLabel.TABLE:
                text = _table_text_from_cells(child)
            else:
                text = (getattr(child, "text", "") or "").strip()
            if text:
                parts.append(text)
            j += 1
        if parts:
            recovered.append((heading, "\n\n".join(parts), i))
    return recovered


def chunk_documents(raw_documents, storage_path=PDF_STORAGE_PATH, classify=False):
    if not raw_documents:
        logger.warning("chunk_documents called with no raw documents.")
        st.warning("No content found in the document to chunk.")
        return [], [], []

    logger.info(f"Starting Docling hybrid chunking on {len(raw_documents)} document(s).")

    try:
        tokenizer = _load_docling_tokenizer()
        if tokenizer is None:
            return [], [], []
        max_tokens = _resolve_max_tokens(tokenizer)
        hf_tokenizer = HuggingFaceTokenizer(tokenizer=tokenizer, max_tokens=max_tokens)
        chunker = HybridChunker(tokenizer=hf_tokenizer, merge_peers=True)

        all_chunks = []
        general_context_chunks = []
        requirements_chunks = []
        processed_chunk_texts = set()
        chunks_pending_classification: list[tuple[str, list, dict]] = []

        docs_by_source: dict[str, list[LangchainDocument]] = {}
        for doc in raw_documents:
            source_path = doc.metadata.get("source")
            if not source_path:
                raise ValueError("Document is missing 'source' metadata.")
            docs_by_source.setdefault(source_path, []).append(doc)

        for source_path, source_docs in docs_by_source.items():
            doc_metadata = source_docs[0].metadata
            normalized_source = _normalize_source_path(source_path)
            is_pdf_source = _is_pdf_source(source_path, source_docs, normalized_source)

            text_blocks = [
                d.page_content for d in source_docs if d.page_content and d.page_content.strip()
            ]
            if not text_blocks:
                logger.warning(f"No extractable text found for '{source_path}'.")
                continue

            doc_name = Path(normalized_source or source_path).stem or "document"

            # Resolve the on-disk path once; it drives both native conversion and
            # (for PDFs) ToC extraction.
            full_path = normalized_source or source_path
            if not os.path.exists(full_path):
                logger.warning(
                    f"Source path '{full_path}' not found. Attempting to resolve with storage path '{storage_path}'."
                )
                full_path = os.path.join(storage_path, os.path.basename(source_path))

            # Prefer Docling's native conversion so SECTION_HEADER structure (and
            # therefore the heading hierarchy that drives the requirements
            # workbook indentation) is preserved. Fall back to a text-only
            # document only when the file isn't on disk or conversion fails.
            dl_doc = None
            if os.path.exists(full_path):
                dl_doc = _convert_with_docling(full_path, is_pdf_source)
            if dl_doc is None:
                dl_doc = _build_docling_document_from_text(text_blocks, doc_name)

            # PDF-only: extract the ToC and build a number->heading index for parent
            # reconstruction, and drop running headers/footers and false headings.
            toc_index = {}
            if is_pdf_source:
                if os.path.exists(full_path):
                    toc_lines = extract_toc_lines(full_path)
                    toc_set = build_toc_set(toc_lines)
                    dl_doc, cleanup_stats = clean_heading_nodes(dl_doc, toc_set)
                    hdrs = [
                        item.text
                        for item, _ in dl_doc.iterate_items()
                        if item.label == DocItemLabel.SECTION_HEADER and item.text.strip()
                    ]
                    toc_index = build_toc_index(hdrs)
                    logger.info(
                        f"PDF '{os.path.basename(full_path)}': cleanup removed "
                        f"{cleanup_stats.get('total_removed', 0)}, {len(toc_lines)} ToC lines "
                        f"-> {len(toc_index)} indexed for parent reconstruction."
                    )
                else:
                    logger.warning(
                        f"PDF path '{full_path}' not found; skipping ToC-based heading expansion."
                    )
                chunks = list(chunker.chunk(dl_doc))
            else:
                chunker = HybridChunker(
                    tokenizer=hf_tokenizer,
                    merge_peers=True,
                    serializer_provider=MDTableSerializerProvider(),   # full markdown tables in chunks
                    # repeat_table_header=True,   # keep header on each split of a wide/long table
                )
                chunks = list(chunker.chunk(dl_doc))

            logger.info(f"Number of chunks before deduplication: {len(chunks)}")

            # Compute a document-order key for every chunk (normal + recovered)
            # so recovered sections are reinserted in the right place instead
            # of appended at the end. `doc_item_seq` is walked once and reused
            # by both the position lookup and the recovery/ancestry helpers.
            # `header_position` maps a header's own (pre-expansion) text to its
            # index -- pre-expansion so it matches the raw item text below even
            # for PDFs, where expand_headings_with_parents may rewrite the leaf.
            doc_item_seq = list(dl_doc.iterate_items())
            header_position: dict[str, int] = {}
            for idx, (doc_item, _) in enumerate(doc_item_seq):
                if doc_item.label == DocItemLabel.SECTION_HEADER:
                    header_position.setdefault(doc_item.text.strip(), idx)
            heading_ancestry = _build_heading_ancestry(doc_item_seq)

            doc_ordered_items: list[tuple[float, int, object]] = []
            tiebreak = 0

            for i, c in enumerate(chunks):
                chunk_text = c.text.strip()
                if not chunk_text or chunk_text in processed_chunk_texts:
                    continue

                processed_chunk_texts.add(chunk_text)

                raw_headings = list(c.meta.headings) if c.meta.headings else []
                raw_leaf = raw_headings[-1] if raw_headings else None
                position = header_position.get(raw_leaf, float("inf"))

                # For PDFs, expand a numbered leaf heading into its full ancestor
                # chain. The ToC index (when PyMuPDF is available) supplies real
                # parent titles; otherwise depth is still reconstructed from the
                # leaf's section number so indentation is preserved. For non-PDFs,
                # the chunker already provides the full heading chain.
                headings = raw_headings
                if is_pdf_source and headings:
                    headings = expand_headings_with_parents(headings, toc_index)

                payload = (
                    (chunk_text, headings, doc_metadata) if classify else
                    LangchainDocument(
                        page_content=chunk_text,
                        metadata={**doc_metadata, "headings": headings, "in_memory": False},
                    )
                )
                doc_ordered_items.append((position, tiebreak, payload))
                tiebreak += 1

            # Recover sections whose body produced no chunk at all -- typically
            # a layout table (no real header row) that the chunker's table
            # serialization rendered as empty text. See
            # _recover_missing_table_sections for why this can't be fixed by
            # export_to_markdown()/export_to_text() alone, and isn't fully
            # covered by MDTableSerializerProvider either (a table can still
            # serialize to nothing in the visited-set edge case). Each
            # recovered section gets its full heading ancestry from
            # _build_heading_ancestry (recovered chunks never go through the
            # chunker's own heading tracking, so without this they'd carry
            # only their own leaf heading and lose their parents).
            for heading, recovered_text, position in _recover_missing_table_sections(
                dl_doc, chunks, doc_item_seq
            ):
                if not recovered_text or recovered_text in processed_chunk_texts:
                    continue
                processed_chunk_texts.add(recovered_text)

                recovered_headings = heading_ancestry.get(position, [heading])
                if is_pdf_source and toc_index:
                    # No-op when the leaf is unnumbered (returns the ancestry
                    # list unchanged); rebuilds the chain from toc_index when
                    # it is, same as normal PDF chunks.
                    recovered_headings = expand_headings_with_parents(recovered_headings, toc_index)

                payload = (
                    (recovered_text, recovered_headings, doc_metadata) if classify else
                    LangchainDocument(
                        page_content=recovered_text,
                        metadata={**doc_metadata, "headings": recovered_headings, "in_memory": False},
                    )
                )
                doc_ordered_items.append((position, tiebreak, payload))
                tiebreak += 1
                logger.info(f"Recovered section skipped by chunker: {heading!r}")

            # Reinsert everything in document order -- recovered sections
            # interleave with their siblings instead of landing at the end.
            doc_ordered_items.sort(key=lambda t: (t[0], t[1]))
            for _, _, payload in doc_ordered_items:
                if classify:
                    chunks_pending_classification.append(payload)
                else:
                    all_chunks.append(payload)

        if classify and chunks_pending_classification:
            language_model = get_language_model()
            if language_model is None:
                logger.error("Language model unavailable; cannot classify document chunks.")
                st.error("Language model unavailable; cannot classify document chunks.")
                return [], [], []

            def _classify_one(item: tuple[str, list, dict]):
                chunk_text, headings, doc_metadata = item
                classification = classify_chunk(language_model, chunk_text)
                return chunk_text, headings, doc_metadata, classification

            workers = min(_CLASSIFY_MAX_WORKERS, len(chunks_pending_classification))
            logger.info(
                f"Classifying {len(chunks_pending_classification)} chunk(s) with {workers} worker(s)."
            )
            with ThreadPoolExecutor(max_workers=workers) as executor:
                classified = list(executor.map(_classify_one, chunks_pending_classification))

            for chunk_text, headings, doc_metadata, classification in classified:
                if classification == "General Context":
                    general_context_chunks.append(
                        LangchainDocument(
                            page_content=chunk_text,
                            metadata={**doc_metadata, "headings": headings, "in_memory": True},
                        )
                    )
                else:
                    requirements_chunks.append(
                        LangchainDocument(
                            page_content=chunk_text,
                            metadata={**doc_metadata, "headings": headings, "in_memory": False},
                        )
                    )

        if classify:
            logger.info(f"Docling hybrid chunking and classification complete.")
            logger.info(f"Number of chunks after deduplication: {len(processed_chunk_texts)}")
            logger.info(f"  - General Context chunks: {len(general_context_chunks)}")
            logger.info(f"  - Requirements chunks: {len(requirements_chunks)}")
        else:
            logger.info(f"Docling hybrid chunking complete: {len(all_chunks)} chunks created.")

        return general_context_chunks, requirements_chunks, all_chunks

    except Exception as e:
        logger.exception(f"An error occurred during hybrid chunking using Docling. Details: {e}")
        st.error("An error occurred during hybrid chunking using Docling. Check logs for details.")
        return [], [], []

def index_documents(document_chunks, vector_db=None):
    if not document_chunks:
        logger.warning("index_documents called with no chunks to index.")
        st.warning("No document chunks available to index.")
        return False
    logger.info(f"Indexing {len(document_chunks)} document chunks.")
    try:
        if vector_db is None:
            vector_db = st.session_state.DOCUMENT_VECTOR_DB
        vector_db.add_documents(document_chunks)
        logger.info("Document chunks indexed successfully into vector store.")
        return True
    except Exception as e:
        user_message = "An error occurred while indexing document chunks."
        logger.exception(f"{user_message} Details: {e}")
        st.error(f"{user_message} Check logs for details.")
        return False


def re_index_documents_from_session():
    """
    Re-indexes documents from chunks stored in the session state.
    This is used to repopulate in-memory vector databases after a session is loaded.
    """
    logger.info("Attempting to re-index documents from session state.")

    # Re-index general context chunks
    if "general_context_chunks" in st.session_state and st.session_state.general_context_chunks:
        logger.info(f"Re-indexing {len(st.session_state.general_context_chunks)} general context chunks.")
        index_documents(st.session_state.general_context_chunks, vector_db=st.session_state.GENERAL_VECTOR_DB)
    else:
        logger.info("No general context chunks found in session state to re-index.")

    # Re-index requirements chunks
    if "requirements_chunks" in st.session_state and st.session_state.requirements_chunks:
        logger.info(f"Re-indexing {len(st.session_state.requirements_chunks)} requirements chunks.")
        index_documents(st.session_state.requirements_chunks, vector_db=st.session_state.DOCUMENT_VECTOR_DB)
    else:
        logger.info("No requirements chunks found in session state to re-index.")

    # Re-index standalone context chunks
    if "context_chunks" in st.session_state and st.session_state.context_chunks:
        logger.info(f"Re-indexing {len(st.session_state.context_chunks)} standalone context chunks.")
        index_documents(st.session_state.context_chunks, vector_db=st.session_state.CONTEXT_VECTOR_DB)
        # Also ensure the loaded flag is set if we are re-indexing its chunks
        st.session_state.context_document_loaded = True
    else:
        logger.info("No standalone context chunks found in session state to re-index.")
