import os
import streamlit as st
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_ollama import OllamaEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.llms import OllamaLLM
from langchain_core.documents import Document as LangchainDocument
import docx # For Document and opc.exceptions
import pdfplumber # For pdfplumber.exceptions
from rank_bm25 import BM25Okapi


# ---------------------------------
# App Styling with CSS
# ---------------------------------
st.markdown(
    """
    <style>
    .stApp {
        background-color: #0E1117;
        color: #FFFFFF;
    }
    
    /* Chat Input Styling */
    .stChatInput input {
        background-color: #1E1E1E !important;
        color: #FFFFFF !important;
        border: 1px solid #3A3A3A !important;
    }
    
    /* User Message Styling */
    .stChatMessage[data-testid="stChatMessage"]:nth-child(odd) {
        background-color: #1E1E1E !important;
        border: 1px solid #3A3A3A !important;
        color: #E0E0E0 !important;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
    
    /* Assistant Message Styling */
    .stChatMessage[data-testid="stChatMessage"]:nth-child(even) {
        background-color: #2A2A2A !important;
        border: 1px solid #404040 !important;
        color: #F0F0F0 !important;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
    
    /* Avatar Styling */
    .stChatMessage .avatar {
        background-color: #00FFAA !important;
        color: #000000 !important;
    }
    
    /* Text Color Fix */
    .stChatMessage p, .stChatMessage div {
        color: #FFFFFF !important;
    }
    
    .stFileUploader {
        background-color: #1E1E1E;
        border: 1px solid #3A3A3A;
        border-radius: 5px;
        padding: 15px;
    }
    
    h1, h2, h3 {
        color: #00FFAA !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------
# Global Configuration
# ---------------------------------
MAX_HISTORY_TURNS = 3 # Number of recent user/assistant turn pairs to include in history
K_SEMANTIC = 5 # Number of results for semantic search
K_BM25 = 5     # Number of results for BM25 search
K_RRF_PARAM = 60 # Constant for Reciprocal Rank Fusion (RRF)

PROMPT_TEMPLATE = """
You are an expert research assistant. Use the provided document context and conversation history to answer the current query.
If the query is a follow-up question, use the conversation history to understand the context.
If unsure, state that you don't know. Be concise and factual (max 3 sentences).

Conversation History (if any):
{conversation_history}

Document Context:
{document_context}

Current Query: {user_query}
Answer:
"""

PDF_STORAGE_PATH = "document_store/pdfs/"

# Ensure the PDF storage directory exists
os.makedirs(PDF_STORAGE_PATH, exist_ok=True)

# Fetch Ollama base URL from environment variable, with a default
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

# Cached functions to load models
@st.cache_resource
def get_embedding_model():
    print(f"Loading Embedding Model from: {OLLAMA_BASE_URL}") # For debugging/verifying cache
    return OllamaEmbeddings(model="deepseek-r1:1.5b", base_url=OLLAMA_BASE_URL)

@st.cache_resource
def get_language_model():
    print(f"Loading Language Model from: {OLLAMA_BASE_URL}") # For debugging/verifying cache
    return OllamaLLM(model="deepseek-r1:1.5b", base_url=OLLAMA_BASE_URL)

# Initialize models using cached functions
EMBEDDING_MODEL = get_embedding_model()
LANGUAGE_MODEL = get_language_model()

# Initialize or retrieve the vector store from session state for persistence across interactions
# Note: The vector store itself is not cached with st.cache_resource here because it's mutable 
# and tied to session state. We re-initialize it in session_state if it's not there or when reset.
# The EMBEDDING_MODEL used by it *is* cached.
if "DOCUMENT_VECTOR_DB" not in st.session_state:
    st.session_state.DOCUMENT_VECTOR_DB = InMemoryVectorStore(EMBEDDING_MODEL)
if "messages" not in st.session_state:
    st.session_state.messages = []
if "document_processed" not in st.session_state:
    st.session_state.document_processed = False
if "uploaded_file_key" not in st.session_state:
    st.session_state.uploaded_file_key = 0
if "raw_documents" not in st.session_state: # To store loaded documents before chunking
    st.session_state.raw_documents = []
if "document_summary" not in st.session_state: # To store the generated summary
    st.session_state.document_summary = None
if "document_keywords" not in st.session_state: # To store extracted keywords
    st.session_state.document_keywords = None
if "bm25_index" not in st.session_state:
    st.session_state.bm25_index = None
if "bm25_corpus_chunks" not in st.session_state:
    st.session_state.bm25_corpus_chunks = []


# ---------------------------------
# Utility Functions
# ---------------------------------
def save_uploaded_file(uploaded_file):
    """
    Save the uploaded file to disk and return the file path.
    """
    file_path = os.path.join(PDF_STORAGE_PATH, uploaded_file.name)
    try:
        with open(file_path, "wb") as file:
            file.write(uploaded_file.getbuffer())
        st.session_state.uploaded_filename = uploaded_file.name # Store filename
        return file_path
    except IOError as e: # More specific for file system issues
        st.error(f"Failed to save uploaded file '{uploaded_file.name}'. An I/O error occurred: {e.strerror}. Please check permissions and disk space.")
        return None
    except Exception as e: # Catch-all for other unexpected errors
        st.error(f"An unexpected error occurred while saving '{uploaded_file.name}': {e}")
        return None

def load_document(file_path):
    """
    Load documents from PDF, DOCX, or TXT files.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)

    try:
        if file_extension == ".pdf":
            try:
                document_loader = PDFPlumberLoader(file_path)
                return document_loader.load()
            except pdfplumber.exceptions.PDFSyntaxError:
                st.error(f"Failed to load PDF '{file_name}': The file may be corrupted or not a valid PDF.")
                return []
            except Exception as e: # Catch other PDFPlumberLoader errors
                st.error(f"Failed to load PDF '{file_name}': An unexpected error occurred during PDF processing. Details: {e}")
                return []
        elif file_extension == ".docx":
            try:
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                if not full_text.strip():
                    st.warning(f"DOCX file '{file_name}' appears to be empty or contains no text.")
                    return []
                return [LangchainDocument(page_content=full_text, metadata={"source": file_name})]
            except docx.opc.exceptions.PackageNotFoundError: # Specific for corrupted/invalid DOCX
                st.error(f"Failed to load DOCX '{file_name}': The file appears to be corrupted or not a valid DOCX file.")
                return []
            except Exception as e: # Catch other python-docx errors
                st.error(f"Failed to load DOCX '{file_name}': An unexpected error occurred. Details: {e}")
                return []
        elif file_extension == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
                if not full_text.strip():
                    st.warning(f"Text file '{file_name}' appears to be empty.")
                    return []
                return [LangchainDocument(page_content=full_text, metadata={"source": file_name})]
            except UnicodeDecodeError:
                st.error(f"Failed to load TXT file '{file_name}': The file is not UTF-8 encoded. Please ensure it's a plain text file with UTF-8 encoding.")
                return []
            except IOError as e:
                st.error(f"Failed to load TXT file '{file_name}': An I/O error occurred. {e.strerror}.")
                return []
            except Exception as e: # Catch other text file errors
                st.error(f"Failed to load TXT file '{file_name}': An unexpected error occurred. Details: {e}")
                return []
        else:
            st.error(f"Unsupported file type: '{file_extension}' for file '{file_name}'. Please upload a PDF, DOCX, or TXT file.")
            return []
    except MemoryError:
        st.error(f"Failed to load document '{file_name}': The file is too large to process with available memory.")
        return []
    except Exception as e: # Catch-all for other unexpected errors during loading dispatch
        st.error(f"An unexpected error occurred while attempting to load '{file_name}': {e}")
        return []

def chunk_documents(raw_documents):
    """
    Split raw documents into manageable chunks using RecursiveCharacterTextSplitter.
    """
    if not raw_documents: # Check if raw_documents is empty
        st.warning("No content found in the document to chunk.")
        return []
    try:
        text_processor = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            add_start_index=True,
        )
        processed_docs = text_processor.split_documents(raw_documents)
        if not processed_docs:
            st.warning("Document chunking resulted in no processable text chunks. The document might be valid but contain no extractable text after initial processing, or the text is too short.")
        return processed_docs
    except Exception as e:
        st.error(f"An error occurred while chunking documents from '{st.session_state.get('uploaded_filename', 'the file')}'. This may be due to unexpected document structure. Details: {e}")
        return []

def index_documents(document_chunks):
    """
    Add document chunks to the in-memory vector store.
    """
    if not document_chunks:
        st.warning("No document chunks available to index. This may happen if the document was empty or text extraction failed.")
        return
    try:
        st.session_state.DOCUMENT_VECTOR_DB.add_documents(document_chunks)
        st.session_state.document_processed = True 
    except Exception as e:
        st.error(f"An error occurred while indexing document chunks for '{st.session_state.get('uploaded_filename', 'the file')}'. Details: {e}")
        st.session_state.document_processed = False

def find_related_documents(query):
    """
    Perform both semantic and BM25 search to find related document chunks.
    Returns a dictionary with 'semantic_results' and 'bm25_results'.
    """
    semantic_docs = []
    bm25_retrieved_chunks = []

    if not query or not query.strip(): # Check for empty query
        st.warning("Search query is empty. Please enter a query to find related documents.")
        return {"semantic_results": semantic_docs, "bm25_results": bm25_retrieved_chunks}

    if not st.session_state.document_processed:
        st.warning("No document has been processed yet. Please upload and process a document before searching.")
        return {"semantic_results": semantic_docs, "bm25_results": bm25_retrieved_chunks}

    # 1. Semantic Search (Vector Search)
    try:
        # Ensure K_SEMANTIC doesn't exceed available docs if the library doesn't handle it
        # similarity_search(k=...) usually handles this, but good to be mindful
        semantic_docs = st.session_state.DOCUMENT_VECTOR_DB.similarity_search(query, k=K_SEMANTIC)
        print(f"Semantic search found {len(semantic_docs)} results.")
    except Exception as e:
        st.error(f"An error occurred during semantic search for '{st.session_state.get('uploaded_filename', 'the file')}'. Details: {e}")
        semantic_docs = [] # Ensure it's an empty list on error

    # 2. BM25 Search
    if st.session_state.get("bm25_index") and st.session_state.get("bm25_corpus_chunks"):
        try:
            tokenized_query = query.lower().split(" ")
            
            all_doc_scores = st.session_state.bm25_index.get_scores(tokenized_query)
            
            # Get indices of top k_bm25 scores
            num_bm25_chunks = len(st.session_state.bm25_corpus_chunks)
            num_docs_to_consider = min(K_BM25, num_bm25_chunks)
            
            # Ensure we only process valid indices and scores
            top_n_indices = sorted(
                [i for i, score in enumerate(all_doc_scores) if score > 0], # Consider only positive scores
                key=lambda i: all_doc_scores[i], 
                reverse=True
            )[:num_docs_to_consider]
            
            bm25_retrieved_chunks = [st.session_state.bm25_corpus_chunks[i] for i in top_n_indices]
            print(f"BM25 search found {len(bm25_retrieved_chunks)} results with positive scores.")

        except Exception as e:
            st.error(f"An error occurred during BM25 search for '{st.session_state.get('uploaded_filename', 'the file')}'. Details: {e}")
            bm25_retrieved_chunks = [] # Ensure it's an empty list on error
    else:
        print("BM25 index not available. Skipping BM25 search.")
        bm25_retrieved_chunks = []

    return {"semantic_results": semantic_docs, "bm25_results": bm25_retrieved_chunks}

def combine_results_rrf(search_results_dict, k_param=K_RRF_PARAM):
    """
    Combines search results from different methods using Reciprocal Rank Fusion (RRF).
    Args:
        search_results_dict (dict): A dictionary where keys are search method names (e.g., "semantic_results", "bm25_results")
                                     and values are lists of Langchain Document objects, ordered by relevance.
        k_param (int): The RRF k parameter (default is K_RRF_PARAM).
    Returns:
        list: A de-duplicated list of Langchain Document objects, sorted by RRF score.
    """
    doc_to_score = {}  # Maps page_content to its RRF score
    doc_objects = {}   # Maps page_content to the actual Document object

    # Process semantic results
    semantic_results = search_results_dict.get("semantic_results", [])
    for i, doc in enumerate(semantic_results):
        doc_id = doc.page_content  # Assuming page_content is a unique identifier for a chunk
        if doc_id not in doc_objects:
            doc_objects[doc_id] = doc
        score = 1.0 / (k_param + i + 1) # Rank is i+1
        doc_to_score[doc_id] = doc_to_score.get(doc_id, 0) + score

    # Process BM25 results
    bm25_results = search_results_dict.get("bm25_results", [])
    for i, doc in enumerate(bm25_results):
        doc_id = doc.page_content
        if doc_id not in doc_objects:
            doc_objects[doc_id] = doc
        score = 1.0 / (k_param + i + 1) # Rank is i+1
        doc_to_score[doc_id] = doc_to_score.get(doc_id, 0) + score
            
    # Sort documents by RRF score in descending order
    sorted_doc_ids = sorted(doc_to_score.keys(), key=lambda x: doc_to_score[x], reverse=True)
    
    # Get the final de-duplicated list of Document objects
    final_combined_docs = [doc_objects[doc_id] for doc_id in sorted_doc_ids]
    
    print(f"RRF combined {len(semantic_results)} semantic and {len(bm25_results)} BM25 results into {len(final_combined_docs)} unique docs.")
    return final_combined_docs


def generate_answer(user_query, context_documents, conversation_history=""):
    """
    Generate an answer based on the user query, context documents, and conversation history.
    """
    if not user_query or not user_query.strip():
        return "Your question is empty. Please type a question to get an answer."
    
    # Check if context_documents is a list and not empty
    if not context_documents or not isinstance(context_documents, list) or len(context_documents) == 0:
        return "I couldn't find relevant information in the document to answer your query. Please try rephrasing your question or ensure the document contains the relevant topics."
        
    try:
        context_text = "\n\n".join([doc.page_content for doc in context_documents])
        if not context_text.strip(): # Edge case: context_documents exist but have no content
             return "The relevant sections found in the document appear to be empty. Cannot generate an answer."
        
        conversation_prompt = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
        response_chain = conversation_prompt | LANGUAGE_MODEL
        
        response = response_chain.invoke({
            "user_query": user_query,
            "document_context": context_text,
            "conversation_history": conversation_history
        })
        
        if not response or not response.strip():
            return "The AI model returned an empty response. Please try rephrasing your question or try again later."
        return response
    except Exception as e:
        st.error(f"An error occurred while generating the answer using the AI model. Details: {e}")
        return "I'm sorry, but I encountered an error while trying to generate a response. Please try again later or rephrase your question."

SUMMARIZATION_PROMPT_TEMPLATE = """
You are an expert research assistant. Provide a concise summary of the following document.
Focus on the main points and key takeaways. The summary should be approximately 3-5 sentences long.

Document:
{document_text}

Summary:
"""

def generate_summary(full_document_text):
    """
    Generates a summary for the given document text.
    """
    if not full_document_text or not full_document_text.strip():
        st.warning("Document content is empty or contains only whitespace. Cannot generate summary.")
        return None # Return None for the UI to handle (e.g. not display summary section)
    try:
        summary_prompt = ChatPromptTemplate.from_template(SUMMARIZATION_PROMPT_TEMPLATE)
        summary_chain = summary_prompt | LANGUAGE_MODEL
        summary = summary_chain.invoke({"document_text": full_document_text})
        if not summary or not summary.strip():
            st.warning("The AI model returned an empty summary. The document might be too short or lack clear content for summarization.")
            return None 
        return summary
    except Exception as e:
        st.error(f"An error occurred while generating the document summary using the AI model. Details: {e}")
        return "Failed to generate summary due to an AI model error. Please try again later."

KEYWORD_EXTRACTION_PROMPT_TEMPLATE = """
You are an expert research assistant. Analyze the following document and extract the top 5-10 most relevant keywords or key phrases.
Present them as a comma-separated list.

Document:
{document_text}

Keywords:
"""

def generate_keywords(full_document_text):
    """
    Generates keywords for the given document text.
    """
    if not full_document_text or not full_document_text.strip():
        st.warning("Document content is empty or contains only whitespace. Cannot extract keywords.")
        return None # Return None for the UI to handle
    try:
        keywords_prompt = ChatPromptTemplate.from_template(KEYWORD_EXTRACTION_PROMPT_TEMPLATE)
        keywords_chain = keywords_prompt | LANGUAGE_MODEL
        keywords = keywords_chain.invoke({"document_text": full_document_text})
        if not keywords or not keywords.strip():
            st.warning("The AI model returned no keywords. The document might be too short or lack distinct terms.")
            return None
        return keywords
    except Exception as e:
        st.error(f"An error occurred while extracting keywords using the AI model. Details: {e}")
        return "Failed to extract keywords due to an AI model error. Please try again later."


# ---------------------------------
# User Interface
# ---------------------------------

# Sidebar for controls
with st.sidebar:
    st.header("Controls")
    if st.button("Clear Chat History", key="clear_chat"):
        st.session_state.messages = []
        # No st.rerun() needed here, Streamlit automatically reruns on widget interaction.

    if st.button("Reset Document", key="reset_document"):
        # Re-initialize with the cached embedding model
        st.session_state.DOCUMENT_VECTOR_DB = InMemoryVectorStore(get_embedding_model())
        st.session_state.document_processed = False
        st.session_state.messages = [] # Clear chat as well, as context is lost
        st.session_state.uploaded_file_key += 1 # Increment key to reset file uploader
        if 'uploaded_filename' in st.session_state:
            del st.session_state.uploaded_filename # Clear stored filename
        st.session_state.raw_documents = [] # Clear raw documents
        st.session_state.document_summary = None # Clear summary
        st.session_state.document_keywords = None # Clear keywords
        st.session_state.bm25_index = None # Clear BM25 index
        st.session_state.bm25_corpus_chunks = [] # Clear BM25 corpus
        st.success("Document and chat reset. You can now upload a new document.")
        st.rerun()

    if st.session_state.document_processed:
        if st.button("Summarize Document", key="summarize_doc_button"):
            if st.session_state.raw_documents:
                with st.spinner("Generating summary... This might take a few moments."):
                    full_text = "\n\n".join([doc.page_content for doc in st.session_state.raw_documents])
                    if not full_text.strip():
                        st.sidebar.warning("Cannot generate summary: The document content is effectively empty.")
                        st.session_state.document_summary = None
                    else:
                        summary_text = generate_summary(full_text)
                        st.session_state.document_summary = summary_text # This might be None or an error message
            else:
                st.sidebar.error("Cannot generate summary: Document content not loaded or available.")
        
        if st.button("Extract Keywords", key="extract_keywords_button"):
            if st.session_state.raw_documents:
                with st.spinner("Extracting keywords..."):
                    full_text = "\n\n".join([doc.page_content for doc in st.session_state.raw_documents])
                    if not full_text.strip():
                        st.sidebar.warning("Cannot extract keywords: The document content is effectively empty.")
                        st.session_state.document_keywords = None
                    else:
                        keywords_text = generate_keywords(full_text)
                        st.session_state.document_keywords = keywords_text # This might be None or an error message
            else:
                st.sidebar.error("Cannot extract keywords: Document content not loaded or available.")
    
    # Display summary if available and not an error message (error messages are shown by st.error within generate_summary)
    if st.session_state.document_summary and "Failed to generate summary" not in st.session_state.document_summary:
        st.sidebar.markdown("---")
        st.sidebar.subheader("📄 Document Summary")
        st.sidebar.info(st.session_state.document_summary)
    elif st.session_state.document_summary and "Failed to generate summary" in st.session_state.document_summary:
        # Error already displayed by st.error in generate_summary, clear it from session state so it doesn't persist as "info"
        # This is a bit of a workaround for functions returning error strings. Ideally, they'd raise exceptions handled by caller.
        pass # Error is already shown by the function

    # Display keywords if available and not an error message
    if st.session_state.document_keywords and "Failed to extract keywords" not in st.session_state.document_keywords:
        st.sidebar.markdown("---")
        st.sidebar.subheader("🔑 Extracted Keywords")
        # Using st.text for keywords as they might be long and st.success has a specific connotation
        st.sidebar.text_area("Keywords:", st.session_state.document_keywords, height=100, disabled=True)
    elif st.session_state.document_keywords and "Failed to extract keywords" in st.session_state.document_keywords:
        pass # Error is already shown by the function

st.title("📘 DocuMind-AI")
st.markdown("### Your Intelligent Document Assistant")
st.markdown("---")

# File Upload Section
uploaded_file = st.file_uploader(
    "Upload Research Document (PDF, DOCX, TXT)",
    type=["pdf", "docx", "txt"],
    help="Select a PDF, DOCX, or TXT document for analysis. The previous document and chat will be reset.",
    accept_multiple_files=False,
    key=f"file_uploader_{st.session_state.uploaded_file_key}" # Use key to allow reset
)

if uploaded_file:
    # If a new file is uploaded, effectively reset the document state
    if not st.session_state.get('uploaded_filename') or st.session_state.uploaded_filename != uploaded_file.name:
        # Re-initialize with the cached embedding model
        st.session_state.DOCUMENT_VECTOR_DB = InMemoryVectorStore(get_embedding_model())
        st.session_state.document_processed = False
        st.session_state.messages = [] # Clear chat as well
        st.session_state.raw_documents = [] # Clear raw documents for new file
        st.session_state.document_summary = None # Clear previous summary
        st.session_state.document_keywords = None # Clear previous keywords
        st.session_state.bm25_index = None # Clear BM25 index for new file
        st.session_state.bm25_corpus_chunks = [] # Clear BM25 corpus for new file
        if 'uploaded_filename' in st.session_state:
             del st.session_state.uploaded_filename

    if not st.session_state.document_processed:
        with st.spinner(f"Processing '{uploaded_file.name}'... This may take a moment."):
            saved_path = save_uploaded_file(uploaded_file)
            if saved_path:
                raw_docs = load_document(saved_path)
                if raw_docs:
                    st.session_state.raw_documents = raw_docs # Store raw documents
                    processed_chunks = chunk_documents(raw_docs)
                    if processed_chunks: 
                        # Vector Indexing (Existing)
                        index_documents(processed_chunks) 

                        # BM25 Indexing (New)
                        if st.session_state.document_processed: # Proceed only if vector indexing was successful
                            try:
                                corpus_texts = [chunk.page_content for chunk in processed_chunks]
                                tokenized_corpus = [doc.lower().split(" ") for doc in corpus_texts]
                                
                                st.session_state.bm25_index = BM25Okapi(tokenized_corpus)
                                st.session_state.bm25_corpus_chunks = processed_chunks
                                print(f"BM25 index created for {st.session_state.get('uploaded_filename', 'document')}")
                                # Success message updated to reflect both types of indexing
                                st.success(f"✅ Document '{uploaded_file.name}' processed and indexed (Vector & BM25) successfully! You can now ask questions or use analysis tools in the sidebar.")
                            except Exception as e:
                                st.error(f"Failed to create BM25 index for '{uploaded_file.name}'. Details: {e}")
                                # Potentially mark BM25 indexing as failed if needed for other logic
                        else:
                            # Error message would have been shown by index_documents or preceding functions
                            st.error(f"Document '{uploaded_file.name}' processed but failed during vector indexing. BM25 indexing skipped.")
                    # else: Error message would have been shown by chunk_documents or load_document
                # else: Error message would have been shown by load_document
            # else: Error message would have been shown by save_uploaded_file

# Display existing chat messages
for message in st.session_state.messages:
    with st.chat_message(message["role"], avatar=message.get("avatar")):
        st.write(message["content"])

# Chat Input Section - only show if a document has been processed
if st.session_state.document_processed:
    user_input = st.chat_input(f"Ask a question about '{st.session_state.get('uploaded_filename', 'the current document')}'...")
    if user_input:
        if not user_input.strip():
            st.warning("Please enter a question.")
        else:
            # Format conversation history before adding the current user input
            # The history should be what was there *before* this new user_input
            # st.session_state.messages already contains the history up to the last assistant response.
            num_messages_to_take = MAX_HISTORY_TURNS * 2 
            chat_log_for_prompt = st.session_state.get("messages", [])[-num_messages_to_take:]

            history_lines = []
            for msg in chat_log_for_prompt:
                role_label = "User" if msg["role"] == "user" else "Assistant"
                history_lines.append(f"{role_label}: {msg['content']}")
            formatted_history = "\n".join(history_lines)

            # Add current user input to messages for display
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.chat_message("user"):
                st.write(user_input)

            with st.spinner("Thinking..."):
                retrieved_results_dict = find_related_documents(user_input) 
                
                # Combine results using RRF
                combined_docs_for_context = combine_results_rrf(retrieved_results_dict)
                
                if not combined_docs_for_context:
                    ai_response = "I could not find relevant sections in the document to answer your question. Please ensure the document contains information related to your query or try rephrasing."
                else:
                    ai_response = generate_answer(
                        user_query=user_input, 
                        context_documents=combined_docs_for_context,
                        conversation_history=formatted_history
                    )
            
            st.session_state.messages.append({"role": "assistant", "content": ai_response, "avatar": "🤖"})
            with st.chat_message("assistant", avatar="🤖"):
                st.write(ai_response) # This will display the response or error/warning string from generate_answer
else:
    st.info("Please upload a PDF, DOCX, or TXT document to begin your session and ask questions.")
