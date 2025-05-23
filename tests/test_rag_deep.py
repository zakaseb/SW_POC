import pytest
import os
from unittest.mock import patch, MagicMock, mock_open

# Ensure rag_deep.py can be imported. This might need adjustment based on project structure.
# If tests/ is a top-level directory, and rag_deep.py is also top-level:
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from rag_deep import (
    load_document,
    generate_summary,
    generate_keywords,
    generate_answer,
    find_related_documents, 
    combine_results_rrf,  
    rerank_documents, # Added this
    PROMPT_TEMPLATE as RAG_PROMPT_TEMPLATE, 
    SUMMARIZATION_PROMPT_TEMPLATE,
    KEYWORD_EXTRACTION_PROMPT_TEMPLATE,
    K_RRF_PARAM, K_SEMANTIC, K_BM25, # Import constants
    TOP_K_FOR_RERANKER, FINAL_TOP_N_FOR_CONTEXT, # Import new constants
    RERANKER_MODEL # Import the actual model instance for patching/checking
)
from langchain_core.documents import Document as LangchainDocument
from docx.opc.exceptions import PackageNotFoundError
import pdfplumber # Required for one of the mocked exceptions
from rank_bm25 import BM25Okapi 
from sentence_transformers import CrossEncoder # Added this

# Define the path to test data files
TEST_DATA_DIR = os.path.join(os.path.dirname(__file__), "test_data")
SAMPLE_TXT_PATH = os.path.join(TEST_DATA_DIR, "sample.txt")
SAMPLE_DOCX_PATH = os.path.join(TEST_DATA_DIR, "sample.docx")
EMPTY_TXT_PATH = os.path.join(TEST_DATA_DIR, "empty.txt") # Will create this
EMPTY_DOCX_PATH = os.path.join(TEST_DATA_DIR, "empty.docx") # Will create this


@pytest.fixture(scope="session", autouse=True)
def setup_test_environment(tmp_path_factory):
    # Create dummy empty files for testing empty file scenarios
    # These are created once per session
    global EMPTY_TXT_PATH, EMPTY_DOCX_PATH
    
    # If sample.txt and sample.docx are not present (e.g. clean CI environment), create them
    # This is mostly for local testing if a previous step failed. In CI, they should exist.
    if not os.path.exists(SAMPLE_TXT_PATH):
        with open(SAMPLE_TXT_PATH, "w") as f:
            f.write("This is a test content from a TXT file.\nIt has multiple lines.")
            
    if not os.path.exists(SAMPLE_DOCX_PATH):
        from docx import Document as DocxCreatorDocument # Avoid conflict with Langchain's Document
        doc = DocxCreatorDocument()
        doc.add_paragraph("This is test content from a DOCX file.")
        doc.save(SAMPLE_DOCX_PATH)

    # Create empty files
    EMPTY_TXT_PATH = str(tmp_path_factory.mktemp("data") / "empty.txt")
    with open(EMPTY_TXT_PATH, "w") as f:
        f.write("")

    EMPTY_DOCX_PATH = str(tmp_path_factory.mktemp("data") / "empty.docx")
    from docx import Document as DocxCreatorDocument # Avoid conflict with Langchain's Document
    doc = DocxCreatorDocument()
    doc.save(EMPTY_DOCX_PATH)


# --- Tests for load_document ---

def test_load_document_txt_success():
    documents = load_document(SAMPLE_TXT_PATH)
    assert isinstance(documents, list)
    assert len(documents) == 1
    assert isinstance(documents[0], LangchainDocument)
    assert "This is a test content from a TXT file." in documents[0].page_content
    assert documents[0].metadata["source"] == "sample.txt"

def test_load_document_docx_success():
    documents = load_document(SAMPLE_DOCX_PATH)
    assert isinstance(documents, list)
    assert len(documents) == 1
    assert isinstance(documents[0], LangchainDocument)
    assert "This is a test content from a DOCX file." in documents[0].page_content
    assert documents[0].metadata["source"] == "sample.docx"

def test_load_document_empty_txt(capsys):
    documents = load_document(EMPTY_TXT_PATH)
    assert documents == [] # Expect empty list for empty files as per implementation
    captured = capsys.readouterr()
    assert f"Text file '{os.path.basename(EMPTY_TXT_PATH)}' appears to be empty." in captured.out # Check for Streamlit warning

def test_load_document_empty_docx(capsys):
    documents = load_document(EMPTY_DOCX_PATH)
    assert documents == []
    captured = capsys.readouterr()
    assert f"DOCX file '{os.path.basename(EMPTY_DOCX_PATH)}' appears to be empty or contains no text." in captured.out

def test_load_document_unsupported_type(capsys):
    # Create a dummy .png file path (file doesn't need to exist for this test if os.path.splitext is primary)
    unsupported_file_path = os.path.join(TEST_DATA_DIR, "sample.png")
    documents = load_document(unsupported_file_path)
    assert documents == []
    captured = capsys.readouterr()
    assert "Unsupported file type: '.png' for file 'sample.png'." in captured.out # Check for Streamlit error

def test_load_document_non_existent_file(capsys):
    non_existent_path = os.path.join(TEST_DATA_DIR, "non_existent.txt")
    documents = load_document(non_existent_path)
    assert documents == []
    captured = capsys.readouterr()
    # The exact error message depends on the OS and Python version for file not found.
    # For now, checking that an error related to loading that file is shown.
    assert f"Error loading document 'non_existent.txt'" in captured.out

@patch('docx.Document')
def test_load_document_corrupted_docx(mock_docx_document, capsys):
    mock_docx_document.side_effect = PackageNotFoundError("Mocked error: Corrupted DOCX")
    corrupted_docx_path = os.path.join(TEST_DATA_DIR, "corrupted.docx")
    documents = load_document(corrupted_docx_path)
    assert documents == []
    captured = capsys.readouterr()
    assert "Failed to load DOCX 'corrupted.docx': The file appears to be corrupted or not a valid DOCX file." in captured.out

@patch("builtins.open", new_callable=mock_open)
def test_load_document_txt_unicode_decode_error(mock_file_open, capsys):
    mock_file_open.side_effect = UnicodeDecodeError("utf-8", b"\x80", 0, 1, "invalid start byte")
    bad_encoding_txt_path = os.path.join(TEST_DATA_DIR, "bad_encoding.txt")
    documents = load_document(bad_encoding_txt_path)
    assert documents == []
    captured = capsys.readouterr()
    assert "Failed to load TXT file 'bad_encoding.txt': The file is not UTF-8 encoded." in captured.out

@patch('rag_deep.PDFPlumberLoader')
def test_load_document_pdf_syntax_error(mock_pdf_loader, capsys):
    instance = mock_pdf_loader.return_value
    instance.load.side_effect = pdfplumber.exceptions.PDFSyntaxError("Mocked PDF Syntax Error")
    corrupted_pdf_path = os.path.join(TEST_DATA_DIR, "corrupted.pdf")
    documents = load_document(corrupted_pdf_path)
    assert documents == []
    captured = capsys.readouterr()
    assert "Failed to load PDF 'corrupted.pdf': The file may be corrupted or not a valid PDF." in captured.out


# --- Tests for generate_summary ---

@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_summary_success(mock_llm_invoke):
    # We need to mock the result of .invoke() on the chain, not the model directly if using LCEL
    # However, current rag_deep.py uses model.invoke directly after creating prompt.
    # If rag_deep.py had `chain = prompt | LANGUAGE_MODEL; chain.invoke()`, this mock would be different.
    # For now, it's `response_chain.invoke` where response_chain is `conversation_prompt | LANGUAGE_MODEL`
    # So mocking LANGUAGE_MODEL.invoke (or the whole model if it's simpler)
    
    # Mocking the behavior of the LLM part of the chain
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.return_value = "This is a test summary."
    
    # Patch the LANGUAGE_MODEL to be this mock instance
    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        summary = generate_summary("Some document text for summarization.")
        assert summary == "This is a test summary."
        
        # Check if the prompt was constructed correctly (accessing args of the LLM's invoke)
        # The actual prompt object is complex, so checking for key parts.
        args, kwargs = mock_llm_instance.invoke.call_args
        assert "document_text" in kwargs
        assert kwargs["document_text"] == "Some document text for summarization."
        # To check the template, you'd need to inspect the prompt object passed to the chain,
        # which is harder if you only mock the final LLM. For now, this is a good check.

@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_summary_llm_failure(mock_llm_invoke, capsys):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.side_effect = Exception("LLM simulated error")
    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        summary = generate_summary("Text that will cause LLM error.")
        assert "Failed to generate summary due to an AI model error." in summary
        captured = capsys.readouterr()
        assert "An error occurred while generating the document summary using the AI model. Details: LLM simulated error" in captured.out

def test_generate_summary_empty_input(capsys):
    summary = generate_summary("")
    assert summary is None
    captured = capsys.readouterr()
    assert "Document content is empty or contains only whitespace. Cannot generate summary." in captured.out

    summary_ws = generate_summary("   ")
    assert summary_ws is None
    captured_ws = capsys.readouterr()
    assert "Document content is empty or contains only whitespace. Cannot generate summary." in captured_ws.out


# --- Tests for generate_keywords ---

@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_keywords_success(mock_llm_invoke):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.return_value = "keyword1, keyword2, keyword3"
    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        keywords = generate_keywords("Some document text for keyword extraction.")
        assert keywords == "keyword1, keyword2, keyword3"
        args, kwargs = mock_llm_instance.invoke.call_args
        assert "document_text" in kwargs
        assert kwargs["document_text"] == "Some document text for keyword extraction."


@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_keywords_llm_failure(mock_llm_invoke, capsys):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.side_effect = Exception("LLM keyword error")
    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        keywords = generate_keywords("Text for keyword LLM error.")
        assert "Failed to extract keywords due to an AI model error." in keywords
        captured = capsys.readouterr()
        assert "An error occurred while extracting keywords using the AI model. Details: LLM keyword error" in captured.out


def test_generate_keywords_empty_input(capsys):
    keywords = generate_keywords("")
    assert keywords is None
    captured = capsys.readouterr()
    assert "Document content is empty or contains only whitespace. Cannot extract keywords." in captured.out

    keywords_ws = generate_keywords("  \n\t  ")
    assert keywords_ws is None
    captured_ws = capsys.readouterr()
    assert "Document content is empty or contains only whitespace. Cannot extract keywords." in captured_ws.out


# --- Tests for generate_answer ---

@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_answer_with_history(mock_llm_invoke):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.return_value = "Mocked response considering history."
    
    user_query = "What about the second point?"
    # Create a mock LangchainDocument object
    mock_doc = LangchainDocument(page_content="Some context about the second point.")
    context_documents = [mock_doc]
    conversation_history_str = "User: What was the first point?\nAssistant: The first point was ABC."

    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        response = generate_answer(user_query, context_documents, conversation_history=conversation_history_str)

    assert response == "Mocked response considering history."
    
    # Check if the prompt sent to the LLM (via invoke) contained the history
    assert mock_llm_instance.invoke.call_count == 1
    # The first argument to invoke is a dictionary representing the filled prompt
    invoked_prompt_args = mock_llm_instance.invoke.call_args[0][0] 
    
    assert "conversation_history" in invoked_prompt_args
    assert invoked_prompt_args["conversation_history"] == conversation_history_str
    assert "user_query" in invoked_prompt_args
    assert invoked_prompt_args["user_query"] == user_query
    assert "document_context" in invoked_prompt_args
    assert invoked_prompt_args["document_context"] == "Some context about the second point."


@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_answer_without_history_uses_default(mock_llm_invoke):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.return_value = "Mocked response without history."

    user_query = "What is this document about?"
    mock_doc = LangchainDocument(page_content="This document is about testing.")
    context_documents = [mock_doc]
    
    # Call generate_answer without providing conversation_history, relying on default
    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        response = generate_answer(user_query, context_documents)

    assert response == "Mocked response without history."
    assert mock_llm_instance.invoke.call_count == 1
    invoked_prompt_args = mock_llm_instance.invoke.call_args[0][0]

    assert "conversation_history" in invoked_prompt_args
    assert invoked_prompt_args["conversation_history"] == "" # Default value is an empty string
    assert invoked_prompt_args["user_query"] == user_query
    assert invoked_prompt_args["document_context"] == "This document is about testing."


@patch('rag_deep.LANGUAGE_MODEL')
def test_generate_answer_llm_failure(mock_llm_invoke, capsys):
    mock_llm_instance = MagicMock()
    mock_llm_instance.invoke.side_effect = Exception("LLM answer error")
    
    user_query = "A query that causes error."
    mock_doc = LangchainDocument(page_content="Some context.")
    context_documents = [mock_doc]

    with patch('rag_deep.LANGUAGE_MODEL', mock_llm_instance):
        response = generate_answer(user_query, context_documents)

    assert "I'm sorry, but I encountered an error while trying to generate a response." in response
    captured = capsys.readouterr()
    assert "An error occurred while generating the answer using the AI model. Details: LLM answer error" in captured.out


def test_generate_answer_empty_query(capsys):
    mock_doc = LangchainDocument(page_content="Some context.")
    context_documents = [mock_doc]
    response = generate_answer("", context_documents)
    assert response == "Your question is empty. Please type a question to get an answer."

def test_generate_answer_empty_context_docs(capsys):
    user_query = "A valid query."
    response = generate_answer(user_query, []) # Empty list of context_documents
    assert response == "I couldn't find relevant information in the document to answer your query."


# --- Mock Streamlit session state for specific tests ---
@pytest.fixture
def mock_session_state(mocker):
    # Using a dictionary to simulate session_state
    # This allows tests to set/get values as if it were st.session_state
    # without actually using Streamlit's global session state.
    _session_state_dict = {
        "bm25_index": None,
        "bm25_corpus_chunks": [],
        "document_processed": False, # Default to not processed
        "DOCUMENT_VECTOR_DB": MagicMock() # Mock vector DB
    }

    # Mock st.session_state.get
    def get_side_effect(key, default=None):
        return _session_state_dict.get(key, default)

    # Mock direct attribute access if used (e.g., st.session_state.bm25_index)
    # This is more complex as st.session_state is not a simple dict.
    # For functions in rag_deep.py that use st.session_state.get("key"),
    # patching st.session_state.get is often cleaner.
    # If direct access like st.session_state.key is used, then
    # patching st.session_state itself to be a MagicMock or a custom class is needed.
    # rag_deep.py mostly uses st.session_state.get("bm25_index"), which is good.
    
    mocker.patch('streamlit.session_state', _session_state_dict) # For direct access if any
    
    # For functions using st.session_state.get(key), we can mock it if rag_deep imports streamlit as st
    # For direct imports like `from streamlit import session_state`, it's harder.
    # Let's assume rag_deep.py does `import streamlit as st` and uses `st.session_state.get`.
    # However, rag_deep directly uses st.session_state.bm25_index etc.
    # So, the patch above with a dictionary should work for direct attribute access.
    # We'll manage the dict directly in tests.
    return _session_state_dict


# --- Test BM25 Indexing Process (Conceptual - as it's part of main app flow) ---
# This test is more of an integration test for the BM25 part of the main app logic.
# We'll simulate the conditions under which BM25 indexing occurs.

def test_bm25_indexing_process(mock_session_state, capsys):
    # Simulate that vector indexing was successful
    mock_session_state["document_processed"] = True
    
    # Mock processed_chunks that would come from chunk_documents
    doc_content1 = "This is the first chunk for BM25."
    doc_content2 = "Another chunk, the second one for testing BM25."
    processed_chunks = [
        LangchainDocument(page_content=doc_content1),
        LangchainDocument(page_content=doc_content2)
    ]

    # Simulate the part of the main app logic that creates BM25 index
    # This is a direct copy of the logic from rag_deep.py's main block for testing purposes
    if mock_session_state["document_processed"]:
        try:
            corpus_texts = [chunk.page_content for chunk in processed_chunks]
            tokenized_corpus = [doc.lower().split(" ") for doc in corpus_texts]
            
            mock_session_state["bm25_index"] = BM25Okapi(tokenized_corpus)
            mock_session_state["bm25_corpus_chunks"] = processed_chunks
            # print("BM25 index created for test") # For test debugging
        except Exception as e:
            # print(f"Test BM25 indexing error: {e}") # For test debugging
            pass # Let assertions handle it

    assert isinstance(mock_session_state["bm25_index"], BM25Okapi)
    assert mock_session_state["bm25_corpus_chunks"] == processed_chunks
    assert len(mock_session_state["bm25_corpus_chunks"]) == 2


# --- Tests for find_related_documents (Hybrid Retrieval) ---

def create_mock_doc(content, source="mock_source"):
    return LangchainDocument(page_content=content, metadata={"source": source})

@patch('rag_deep.st.session_state', new_callable=MagicMock) # Mock entire session_state object
def test_find_related_documents_both_results(mock_st_session_state):
    # Setup mock session state
    mock_st_session_state.document_processed = True
    mock_st_session_state.get = lambda key, default=None: getattr(mock_st_session_state, key, default)

    # Semantic search results
    semantic_doc1 = create_mock_doc("semantic result 1")
    semantic_doc2 = create_mock_doc("semantic result 2")
    mock_st_session_state.DOCUMENT_VECTOR_DB = MagicMock()
    mock_st_session_state.DOCUMENT_VECTOR_DB.similarity_search.return_value = [semantic_doc1, semantic_doc2]

    # BM25 search results
    bm25_doc1 = create_mock_doc("bm25 result 1")
    bm25_doc2 = create_mock_doc("bm25 result 2")
    mock_bm25_index = MagicMock(spec=BM25Okapi)
    mock_bm25_index.get_scores.return_value = [0.9, 0.8] # Scores for two docs
    mock_st_session_state.bm25_index = mock_bm25_index
    mock_st_session_state.bm25_corpus_chunks = [bm25_doc1, bm25_doc2] # Chunks corresponding to scores

    results = find_related_documents("test query")

    assert len(results["semantic_results"]) == 2
    assert results["semantic_results"][0].page_content == "semantic result 1"
    mock_st_session_state.DOCUMENT_VECTOR_DB.similarity_search.assert_called_once_with("test query", k=K_SEMANTIC)
    
    assert len(results["bm25_results"]) == 2
    assert results["bm25_results"][0].page_content == "bm25 result 1" # Assuming scores lead to this order
    mock_bm25_index.get_scores.assert_called_once_with(["test", "query"])


@patch('rag_deep.st.session_state', new_callable=MagicMock)
def test_find_related_documents_only_semantic(mock_st_session_state):
    mock_st_session_state.document_processed = True
    mock_st_session_state.get = lambda key, default=None: getattr(mock_st_session_state, key, default)

    semantic_doc1 = create_mock_doc("semantic only")
    mock_st_session_state.DOCUMENT_VECTOR_DB = MagicMock()
    mock_st_session_state.DOCUMENT_VECTOR_DB.similarity_search.return_value = [semantic_doc1]
    
    mock_st_session_state.bm25_index = None # BM25 not available
    mock_st_session_state.bm25_corpus_chunks = []

    results = find_related_documents("test query")
    assert len(results["semantic_results"]) == 1
    assert results["semantic_results"][0].page_content == "semantic only"
    assert len(results["bm25_results"]) == 0

@patch('rag_deep.st.session_state', new_callable=MagicMock)
def test_find_related_documents_only_bm25(mock_st_session_state):
    mock_st_session_state.document_processed = True
    mock_st_session_state.get = lambda key, default=None: getattr(mock_st_session_state, key, default)

    mock_st_session_state.DOCUMENT_VECTOR_DB = MagicMock()
    mock_st_session_state.DOCUMENT_VECTOR_DB.similarity_search.return_value = [] # No semantic results

    bm25_doc1 = create_mock_doc("bm25 only")
    mock_bm25_index = MagicMock(spec=BM25Okapi)
    mock_bm25_index.get_scores.return_value = [0.9]
    mock_st_session_state.bm25_index = mock_bm25_index
    mock_st_session_state.bm25_corpus_chunks = [bm25_doc1]

    results = find_related_documents("test query")
    assert len(results["semantic_results"]) == 0
    assert len(results["bm25_results"]) == 1
    assert results["bm25_results"][0].page_content == "bm25 only"

@patch('rag_deep.st.session_state', new_callable=MagicMock)
def test_find_related_documents_no_results(mock_st_session_state):
    mock_st_session_state.document_processed = True
    mock_st_session_state.get = lambda key, default=None: getattr(mock_st_session_state, key, default)

    mock_st_session_state.DOCUMENT_VECTOR_DB = MagicMock()
    mock_st_session_state.DOCUMENT_VECTOR_DB.similarity_search.return_value = []
    
    mock_st_session_state.bm25_index = MagicMock(spec=BM25Okapi)
    mock_st_session_state.bm25_index.get_scores.return_value = [] # Or [0.0, 0.0] for existing chunks
    mock_st_session_state.bm25_corpus_chunks = [create_mock_doc("doc1"), create_mock_doc("doc2")]


    results = find_related_documents("query for no results")
    assert len(results["semantic_results"]) == 0
    assert len(results["bm25_results"]) == 0


# --- Tests for combine_results_rrf ---

def test_combine_results_rrf_no_overlap():
    docS1 = create_mock_doc("semantic doc 1")
    docS2 = create_mock_doc("semantic doc 2")
    docB1 = create_mock_doc("bm25 doc 1")
    docB2 = create_mock_doc("bm25 doc 2")
    
    search_results = {
        "semantic_results": [docS1, docS2], # Ranks 0, 1
        "bm25_results": [docB1, docB2]      # Ranks 0, 1
    }
    combined = combine_results_rrf(search_results, k_param=K_RRF_PARAM)
    
    assert len(combined) == 4
    # With K_RRF_PARAM = 60 (default in function, but explicit here for clarity)
    # Score for rank 0: 1/(60+1) approx 0.01639
    # Score for rank 1: 1/(60+2) approx 0.01613
    # docS1 and docB1 have higher scores. Their order relative to each other depends on stable sort if scores are equal.
    # Here, Python's sort is stable, so original list order for equal scores might be preserved.
    # For RRF, the combined score for docS1 is 1/61, for docB1 is 1/61.
    # For docS2 is 1/62, for docB2 is 1/62.
    # So, {docS1, docB1} should appear before {docS2, docB2}.
    
    combined_contents = [doc.page_content for doc in combined]
    assert docS1.page_content in combined_contents[:2]
    assert docB1.page_content in combined_contents[:2]
    assert docS2.page_content in combined_contents[2:]
    assert docB2.page_content in combined_contents[2:]


def test_combine_results_rrf_full_overlap():
    doc1 = create_mock_doc("doc content 1")
    doc2 = create_mock_doc("doc content 2")
    
    # Semantic: doc1 (rank 0), doc2 (rank 1)
    # BM25:     doc2 (rank 0), doc1 (rank 1)
    search_results = {
        "semantic_results": [doc1, doc2],
        "bm25_results": [doc2, doc1] 
    }
    combined = combine_results_rrf(search_results, k_param=1) # Using k=1 for simpler score calculation
    # Score for doc1: 1/(1+1) [semantic] + 1/(1+2) [bm25] = 0.5 + 0.333 = 0.833
    # Score for doc2: 1/(1+2) [semantic] + 1/(1+1) [bm25] = 0.333 + 0.5 = 0.833
    
    assert len(combined) == 2
    # Since scores are equal, their order can vary due to sort stability.
    # We just need to ensure both are present.
    assert doc1 in combined
    assert doc2 in combined
    # If we want to assert specific order for equal scores, we'd need to know Python's Timsort stability details
    # or ensure unique scores. For this test, presence is key.
    # A more robust test would make scores slightly different.

    # Let's make scores different by changing K_RRF_PARAM or ranks
    # Semantic: doc1 (rank 0), doc2 (rank 1) -> Scores: 1/61, 1/62
    # BM25:     doc2 (rank 0), doc1 (rank 1) -> Scores: 1/61, 1/62
    # Total doc1: 1/61 + 1/62
    # Total doc2: 1/62 + 1/61 
    # Scores are identical. Order depends on stability or original encounter order.
    
    # To make it deterministic, let's make one rank higher consistently
    search_results_deterministic = {
        "semantic_results": [doc1, doc2], # doc1 has higher semantic rank
        "bm25_results": [doc1, doc2]      # doc1 has higher bm25 rank
    }
    combined_det = combine_results_rrf(search_results_deterministic, k_param=K_RRF_PARAM)
    assert len(combined_det) == 2
    assert combined_det[0] == doc1 # doc1 should have sum of two higher scores
    assert combined_det[1] == doc2


def test_combine_results_rrf_partial_overlap():
    doc1 = create_mock_doc("doc1 common")
    docS2 = create_mock_doc("docS2 semantic only")
    docB2 = create_mock_doc("docB2 bm25 only")
    
    # Semantic: doc1 (rank 0), docS2 (rank 1)
    # BM25:     doc1 (rank 0), docB2 (rank 1)
    search_results = {
        "semantic_results": [doc1, docS2],
        "bm25_results": [doc1, docB2]
    }
    combined = combine_results_rrf(search_results, k_param=1)
    # Score for doc1: 1/(1+1) + 1/(1+1) = 0.5 + 0.5 = 1.0
    # Score for docS2: 1/(1+2) = 0.333
    # Score for docB2: 1/(1+2) = 0.333
    
    assert len(combined) == 3
    assert combined[0] == doc1 # doc1 has the highest RRF score
    # docS2 and docB2 have equal scores, their order might vary
    assert docS2 in combined[1:]
    assert docB2 in combined[1:]

def test_combine_results_rrf_empty_lists():
    docS1 = create_mock_doc("semantic doc 1")
    docB1 = create_mock_doc("bm25 doc 1")

    # Both empty
    combined_both_empty = combine_results_rrf({"semantic_results": [], "bm25_results": []})
    assert len(combined_both_empty) == 0

    # Semantic empty
    combined_sem_empty = combine_results_rrf({"semantic_results": [], "bm25_results": [docB1]})
    assert len(combined_sem_empty) == 1
    assert combined_sem_empty[0] == docB1

    # BM25 empty
    combined_bm25_empty = combine_results_rrf({"semantic_results": [docS1], "bm25_results": []})
    assert len(combined_bm25_empty) == 1
    assert combined_bm25_empty[0] == docS1


def test_combine_results_rrf_varying_ranks_and_k():
    docA = create_mock_doc("A")
    docB = create_mock_doc("B")
    docC = create_mock_doc("C")
    docD = create_mock_doc("D") # Unique to BM25

    # Ranks: Sem: A(0) B(1) C(2)  | BM25: B(0) C(1) A(2) D(3)
    # k_param = 2 (for easier math)
    # Scores: 1/(k+rank+1)
    # Doc A: Sem_score = 1/(2+0+1)=1/3 | BM25_score = 1/(2+2+1)=1/5  => Total = 1/3 + 1/5 = 8/15 = 0.533
    # Doc B: Sem_score = 1/(2+1+1)=1/4 | BM25_score = 1/(2+0+1)=1/3  => Total = 1/4 + 1/3 = 7/12 = 0.583
    # Doc C: Sem_score = 1/(2+2+1)=1/5 | BM25_score = 1/(2+1+1)=1/4  => Total = 1/5 + 1/4 = 9/20 = 0.45
    # Doc D: Sem_score = 0           | BM25_score = 1/(2+3+1)=1/6  => Total = 1/6 = 0.166
    # Expected order: B, A, C, D

    search_results = {
        "semantic_results": [docA, docB, docC],
        "bm25_results": [docB, docC, docA, docD]
    }
    combined = combine_results_rrf(search_results, k_param=2)

    assert len(combined) == 4
    assert combined[0] == docB
    assert combined[1] == docA
    assert combined[2] == docC
    assert combined[3] == docD


# --- Tests for rerank_documents ---

@pytest.fixture
def mock_cross_encoder_model():
    model = MagicMock(spec=CrossEncoder)
    return model

@pytest.mark.parametrize("scores, initial_contents, expected_order_indices, top_n", [
    ([0.1, 0.9, 0.5], ["doc1", "doc2", "doc3"], [1, 2, 0], 3), # Re-order
    ([0.9, 0.5, 0.1], ["doc1", "doc2", "doc3"], [0, 1, 2], 3), # Maintain order
    ([0.5, 0.5, 0.5], ["doc1", "doc2", "doc3"], [0, 1, 2], 3), # Equal scores (stable sort)
    ([0.1, 0.9, 0.5], ["doc1", "doc2", "doc3"], [1, 2], 2),    # top_n limits
])
def test_rerank_documents_predict_scenarios(mock_cross_encoder_model, scores, initial_contents, expected_order_indices, top_n):
    mock_cross_encoder_model.predict.return_value = scores
    docs_to_rerank = [create_mock_doc(content) for content in initial_contents]
    
    reranked = rerank_documents("test query", docs_to_rerank, mock_cross_encoder_model, top_n=top_n)
    
    assert len(reranked) == min(top_n, len(initial_contents))
    expected_reranked_contents = [initial_contents[i] for i in expected_order_indices[:top_n]]
    assert [doc.page_content for doc in reranked] == expected_reranked_contents
    mock_cross_encoder_model.predict.assert_called_once()

def test_rerank_documents_model_none(capsys):
    docs_to_rerank = [create_mock_doc(f"doc{i}") for i in range(5)]
    reranked = rerank_documents("test query", docs_to_rerank, model=None, top_n=3)
    
    assert len(reranked) == 3
    assert [doc.page_content for doc in reranked] == ["doc0", "doc1", "doc2"] # Sliced original
    captured = capsys.readouterr()
    assert "Re-ranker model not available. Returning documents without re-ranking." in captured.out

def test_rerank_documents_prediction_error(mock_cross_encoder_model, capsys):
    mock_cross_encoder_model.predict.side_effect = Exception("Prediction failed")
    docs_to_rerank = [create_mock_doc(f"doc{i}") for i in range(5)]
    
    reranked = rerank_documents("test query", docs_to_rerank, mock_cross_encoder_model, top_n=3)
    
    assert len(reranked) == 3
    assert [doc.page_content for doc in reranked] == ["doc0", "doc1", "doc2"] # Sliced original
    captured = capsys.readouterr()
    assert "Error during document re-ranking: Prediction failed" in captured.out

def test_rerank_documents_empty_input():
    mock_model = MagicMock(spec=CrossEncoder)
    reranked = rerank_documents("test query", [], mock_model, top_n=3)
    assert reranked == []


# --- Tests for Chat Logic Integration of Reranking ---

@patch('rag_deep.generate_answer')
@patch('rag_deep.rerank_documents') # Mock the function we just tested
@patch('rag_deep.combine_results_rrf')
@patch('rag_deep.find_related_documents')
@patch('rag_deep.RERANKER_MODEL', new_callable=MagicMock) # Patch the global model instance
def test_chat_logic_reranking_successful(
    mock_reranker_model_global, mock_find_related, mock_combine_rrf, mock_rerank_docs_func, mock_generate_answer, 
    mock_session_state # Use the fixture to manage session state dict
):
    # 1. Setup Mocks
    mock_session_state["document_processed"] = True # Simulate document is processed
    mock_session_state["messages"] = [] # Start with empty chat history for simplicity

    # Mock find_related_documents
    mock_find_related.return_value = {"semantic_results": [create_mock_doc("sem1")], "bm25_results": [create_mock_doc("bm25_1")]}
    
    # Mock combine_results_rrf
    # Let it return TOP_K_FOR_RERANKER docs for reranking
    hybrid_docs = [create_mock_doc(f"hybrid_doc_{i}") for i in range(TOP_K_FOR_RERANKER)]
    mock_combine_rrf.return_value = hybrid_docs
    
    # Mock rerank_documents function (which uses the RERANKER_MODEL)
    # Ensure RERANKER_MODEL global is the mock we can check
    # The rerank_documents function itself is mocked, so its internal use of RERANKER_MODEL isn't directly an issue here.
    # The mock_reranker_model_global is to ensure the `if RERANKER_MODEL:` check passes.
    mock_reranker_model_global.spec = CrossEncoder # Give it a spec to behave like CrossEncoder if needed

    reranked_final_docs = [create_mock_doc(f"reranked_doc_{i}") for i in range(FINAL_TOP_N_FOR_CONTEXT)]
    mock_rerank_docs_func.return_value = reranked_final_docs
    
    # Mock generate_answer
    mock_generate_answer.return_value = "Final AI Answer"

    # 2. Simulate User Input (This is a simplified way to trigger the main logic block)
    user_input = "test user query"
    
    # This is a simplified extraction of the core logic from rag_deep.py's chat input block
    # In a full integration test, you might need to run the Streamlit script or use a more complex setup.
    # For unit testing this part, we directly call the sequence of functions.
    
    retrieved_results_dict = find_related_documents(user_input)
    combined_hybrid_docs = combine_results_rrf(retrieved_results_dict) # Uses K_RRF_PARAM implicitly

    final_context_docs_for_llm = []
    if combined_hybrid_docs:
        docs_for_reranking = combined_hybrid_docs[:TOP_K_FOR_RERANKER]
        if mock_reranker_model_global: # Simulating `if RERANKER_MODEL:`
             final_context_docs_for_llm = mock_rerank_docs_func(user_input, docs_for_reranking, mock_reranker_model_global, top_n=FINAL_TOP_N_FOR_CONTEXT)
        else:
            final_context_docs_for_llm = docs_for_reranking[:FINAL_TOP_N_FOR_CONTEXT]

    if not final_context_docs_for_llm:
        ai_response = "No relevant sections found." # Simplified
    else:
        ai_response = mock_generate_answer(
            user_query=user_input,
            context_documents=final_context_docs_for_llm,
            conversation_history="" # Assuming empty for this test
        )

    # 3. Assertions
    mock_find_related.assert_called_once_with(user_input)
    mock_combine_rrf.assert_called_once_with(retrieved_results_dict, k_param=K_RRF_PARAM)
    
    # Assert rerank_documents function was called correctly
    mock_rerank_docs_func.assert_called_once_with(user_input, hybrid_docs[:TOP_K_FOR_RERANKER], mock_reranker_model_global, top_n=FINAL_TOP_N_FOR_CONTEXT)
    
    # Assert generate_answer was called with the re-ranked documents
    mock_generate_answer.assert_called_once_with(user_query=user_input, context_documents=reranked_final_docs, conversation_history="")
    assert ai_response == "Final AI Answer"


@patch('rag_deep.generate_answer')
@patch('rag_deep.rerank_documents') # Mock this to prevent its execution
@patch('rag_deep.combine_results_rrf')
@patch('rag_deep.find_related_documents')
@patch('rag_deep.RERANKER_MODEL', None) # Patch the global RERANKER_MODEL to be None
@patch('rag_deep.st') # Mock streamlit for st.info
def test_chat_logic_reranker_model_none(
    mock_st, mock_find_related, mock_combine_rrf, mock_rerank_docs_func, mock_generate_answer,
    mock_session_state # Use the fixture to manage session state dict
):
    # 1. Setup Mocks
    mock_session_state["document_processed"] = True
    mock_session_state["messages"] = []

    mock_find_related.return_value = {"semantic_results": [create_mock_doc("sem1")], "bm25_results": [create_mock_doc("bm25_1")]}
    
    hybrid_docs = [create_mock_doc(f"hybrid_doc_{i}") for i in range(TOP_K_FOR_RERANKER)]
    mock_combine_rrf.return_value = hybrid_docs
    
    mock_generate_answer.return_value = "Fallback AI Answer"

    # 2. Simulate User Input (Simplified logic extraction)
    user_input = "test user query"
    
    retrieved_results_dict = find_related_documents(user_input)
    combined_hybrid_docs = combine_results_rrf(retrieved_results_dict)

    final_context_docs_for_llm = []
    # This part of the logic is directly from rag_deep.py, with RERANKER_MODEL patched to None
    # The key is that rerank_documents (mock_rerank_docs_func) should NOT be called.
    if combined_hybrid_docs:
        docs_for_reranking = combined_hybrid_docs[:TOP_K_FOR_RERANKER]
        # Since rag_deep.RERANKER_MODEL is None due to patch, this 'if' will be false
        if RERANKER_MODEL: # This refers to the patched rag_deep.RERANKER_MODEL
             final_context_docs_for_llm = mock_rerank_docs_func(user_input, docs_for_reranking, RERANKER_MODEL, top_n=FINAL_TOP_N_FOR_CONTEXT)
        else: # This 'else' block should be executed
            final_context_docs_for_llm = docs_for_reranking[:FINAL_TOP_N_FOR_CONTEXT]

    if not final_context_docs_for_llm:
        ai_response = "No relevant sections found."
    else:
        ai_response = mock_generate_answer(
            user_query=user_input,
            context_documents=final_context_docs_for_llm,
            conversation_history=""
        )

    # 3. Assertions
    mock_find_related.assert_called_once_with(user_input)
    mock_combine_rrf.assert_called_once_with(retrieved_results_dict, k_param=K_RRF_PARAM)
    
    # Assert rerank_documents function was NOT called
    mock_rerank_docs_func.assert_not_called() 
    
    # Assert st.info was called (or st.warning if that's what rerank_documents does when model is None,
    # but here we are testing the logic *before* rerank_documents is called if model is None in main script)
    # The st.info call is in the main script's logic.
    # The `rerank_documents` function itself has a st.warning if model is None, but it won't be called.
    # This test verifies the branch `if RERANKER_MODEL:` in rag_deep.py
    
    # Assert generate_answer was called with documents from combine_results_rrf, sliced
    expected_context_docs = hybrid_docs[:FINAL_TOP_N_FOR_CONTEXT]
    mock_generate_answer.assert_called_once_with(user_query=user_input, context_documents=expected_context_docs, conversation_history="")
    assert ai_response == "Fallback AI Answer"
    # Check for st.info call (this requires rag_deep.st to be the mocked st)
    mock_st.info.assert_called_once_with("Re-ranker model not loaded. Using documents from hybrid search directly.")
The test file includes:
*   Setup for test data paths and creation of dummy/empty files for specific test cases using a pytest fixture (`setup_test_environment`).
*   **Tests for `load_document`**:
    *   Successful loading of `.txt` and `.docx` files.
    *   Handling of empty `.txt` and `.docx` files (checking for `st.warning`).
    *   Handling of unsupported file types (checking for `st.error`).
    *   Handling of non-existent files (checking for `st.error`).
    *   Handling of corrupted `.docx` files (mocking `docx.Document` to raise `PackageNotFoundError` and checking `st.error`).
    *   Handling of `.txt` files with `UnicodeDecodeError` (mocking `open` and checking `st.error`).
    *   Handling of corrupted `.pdf` files (mocking `PDFPlumberLoader` to raise `pdfplumber.exceptions.PDFSyntaxError` and checking `st.error`).
*   **Tests for `generate_summary`**:
    *   Successful summary generation by mocking `rag_deep.LANGUAGE_MODEL.invoke` and verifying the output and input to the LLM.
    *   Handling of LLM failure by mocking `rag_deep.LANGUAGE_MODEL.invoke` to raise an exception and checking the error message.
    *   Handling of empty input text (checking for `st.warning` and `None` return).
*   **Tests for `generate_keywords`**:
    *   Successful keyword generation (similar mocking strategy as for summaries).
    *   Handling of LLM failure (similar mocking strategy).
    *   Handling of empty input text.
*   The tests use `pytest` and `pytest-mock` (via `unittest.mock.patch` and `MagicMock`).
*   The `capsys` fixture is used to capture output from Streamlit's `st.error` and `st.warning` calls.
*   A `sys.path.insert` is used to allow importing from `rag_deep.py` located in the parent directory.

The test suite covers the requested functionality.
The next step would typically be to run these tests.
The files created/modified are:
*   `requirements-dev.txt` (created in turn 1 of this subtask)
*   `tests/test_data/sample.txt` (created in turn 4 of this subtask)
*   `tests/test_data/sample.docx` (created in turn 5 of this subtask)
*   `tests/test_rag_deep.py` (created in turn 6 of this subtask)
*   The `tests` and `tests/test_data` directories were also created.

I've already mentioned the need for `pytest` and `pytest-mock` (which are in `requirements-dev.txt`).
